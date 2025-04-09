#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Apr  9 11:39:59 2025

@author: mliu
"""


import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
from docx import Document

def bad_rate_by_var_to_doc(dev_df, cols_to_include,r_var, output_dir, max_length=20):
#export tables for all used variables to a work file with each table formatted and saved as a plot
#dev_df: the dataframe where the bad rate by variable is calculated
#cols_to_include: variables included for the calculation
#r_var: response variable, such as 'bk_18'
#output_dir: output directory for saving the docx file
#max_length: wrapping the variable name when it exceeds max_length, defaulted at 20

    doc = Document()
    doc.add_heading('Bad Rate by Variables', 0)
    
    # Function to wrap text in the column names
    def wrap_column_names(columns, max_length=20):
        wrapped_columns = []
        for col in columns:
            # Wrap the text into multiple lines if it's longer than max_length
            wrapped_col = '\n'.join([col[i:i+max_length] for i in range(0, len(col), max_length)])
            wrapped_columns.append(wrapped_col)
        return wrapped_columns
    
    for var in cols_to_include:    
    #    var='s114s'
        df_var=dev_df[[r_var,var]]
        
        if df_var[var].min()<-20:
            df_var['bin_labels'], bin_edges = pd.qcut(df_var[var], q=10, labels=False, retbins=True, duplicates='drop')
            bin_edges_rounded = np.round(bin_edges, 3)
            df_var[var] = df_var['bin_labels'].apply(lambda x: f"({bin_edges_rounded[x]} , {bin_edges_rounded[x+1]}]")
            df=df_var
            
        elif df_var[var].nunique()<=12:
            df_var['bin_labels'] = df_var[var].apply(lambda x: f"{x}")
            df_var[var] = df_var['bin_labels']
            df=df_var
            
        else:    
            positive_values = df_var[df_var[var] > 0]
            non_positive_values = df_var[df_var[var] <= 0]
            
            # Assign separate bins for each non-positive value
            non_positive_values['bin_labels'] = non_positive_values[var].apply(lambda x: f"{x}")  # Each non-positive value gets its own bin
            non_positive_values[var] = non_positive_values['bin_labels']  # For non-positive values, bin_range is the value itself
            
            positive_values['bin_labels'], bin_edges = pd.qcut(positive_values[var], q=10, labels=False, retbins=True, duplicates='drop')
            
            #if data too skewed, increase the number of bins
            if positive_values['bin_labels'].nunique()<5:
                positive_values['bin_labels'], bin_edges = pd.qcut(positive_values[var], q=20, labels=False, retbins=True, duplicates='drop')
            
            # Add a new column for the range of each bin in the format (low, high]
            bin_edges_rounded = np.round(bin_edges, 3)
            bin_edges_rounded[0]=0
            
            #accomodate binary variables that can not be labeled with a range
            if positive_values[var].nunique()>2:
                positive_values[var] = positive_values['bin_labels'].apply(lambda x: f"({bin_edges_rounded[x]} , {bin_edges_rounded[x+1]}]")
                positive_values['bin_labels']=positive_values['bin_labels']+1
            else:
                positive_values['bin_labels'] = positive_values[var].apply(lambda x: f"{x}")  # Each non-positive value gets its own bin
                positive_values[var] = positive_values['bin_labels']  # For non-positive values, bin_range is the value itself            
    
            
            df = pd.concat([non_positive_values,positive_values])
        
        count_by_bin=df[[var,'bin_labels']].value_counts()
        sum_by_bin = df.groupby([var,'bin_labels'])[r_var].sum()
        
        df_sum=pd.concat([count_by_bin, sum_by_bin], axis=1).reset_index(level=var)
        df_sum = df_sum.rename(columns={'count': 'Records', r_var: 'Bads'})
        df_sum['Bad Rate']=(df_sum['Bads']/df_sum['Records'])* 100  # Multiply by 100 to get percentage
        df_sum['Bad Rate'] = df_sum['Bad Rate'].apply(lambda x: f"{x:.2f}%")
        df_sum.index = df_sum.index.astype(float).astype(int)
        df_sum=df_sum.sort_index()
        if df_sum[var].iloc[0].startswith('('):
            df_sum[var].iloc[0] = '[' + df_sum[var].iloc[0][1:]
        
        
            # Create a figure and axis for plotting
        num_rows, num_cols = df_sum.shape
        fig_width = num_cols *1 # Width of the figure (adjust as needed)
        fig_height = num_rows * 0.2  # Height of the figure (adjust as needed)
        
        # Create a figure and axis with dynamic size
        wrapped_columns = wrap_column_names(df_sum.columns, max_length=max_length)
        fig, ax = plt.subplots(figsize=(fig_width, fig_height))
        
        # Hide the axes
        ax.axis('tight')
        ax.axis('off')
        table = ax.table(cellText=df_sum.values, colLabels=wrapped_columns, loc='center')
        # Automatically adjust column widths
        for i in range(len(df_sum.columns)):  # Loop over each column
            max_width = max(df_sum.iloc[:, i].apply(lambda x: len(str(x))))  # Find max length in each column
            table.auto_set_column_width([i])  # Adjust column width based on content
            table.get_celld()[(0, i)].set_width(max_width * 0.1)  # Manually scale width
    
    
    # Adjust the header row height for wrapped column names\
        header_lines = wrapped_columns[0].count('\n') + 1  # Count lines in the header
    
        if header_lines>1:    
            for j in range(len(df_sum.columns)):
                cell = table[(0, j)]
                cell.set_height(0.1 * header_lines) 
    
        table.auto_set_font_size(False)
        table.set_fontsize(11)  # Set font size for the table text
        table.scale(1.5, 1.5)  # Scale the table for better visibility
        # Save the figure as a PNG image
        plt.savefig("table_image.png", bbox_inches="tight", pad_inches=0,dpi=500)
        plt.close()  # Close the plot to avoid overlapping images
    
        doc.add_paragraph(var)
        doc.add_picture("table_image.png")  # Adjust the image width
    
    # Add the DataFrame as a table to the Word document
    #    table = doc.add_table(rows=1, cols=len(df_sum.columns) + len(df_sum.index.names))
    
        # Adding headers to the table
    #    hdr_cells = table.rows[0].cells
    #    for i, column in enumerate(df_sum.columns):
    #        hdr_cells[i].text = column
        
        # Add rows of data to the table (excluding the index)
    #    for idx, row in df_sum.iterrows():
    #        row_cells = table.add_row().cells
            # Add the actual column values (skip the index)
    #        for j, value in enumerate(row):
    #            row_cells[j].text = str(value)
    
    
    # Save the document
    doc.save(output_dir+'bad_by_var.docx')